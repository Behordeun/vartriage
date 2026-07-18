"""Report template engine for clinical variant reports.

Renders assembled report sections into self-contained HTML, PDF
(via WeasyPrint), or DOCX (via python-docx). All output is produced
without network access or external dependencies beyond the optional
rendering libraries.
"""

from __future__ import annotations

import html
from pathlib import Path
from typing import Any

from vartriage.reporting.clinical.models import (EvidenceCardData,
                                                 ExecutiveSummaryData,
                                                 FindingsRow, HeaderData,
                                                 MethodologyData,
                                                 ReportSections, SignOffData)
from vartriage.reporting.clinical.templates import (
    EVIDENCE_CARD_AF_LINE, EVIDENCE_CARD_CLINVAR_LINE,
    EVIDENCE_CARD_INHERITANCE_LINE, EVIDENCE_CARD_SCORES_LINE,
    EVIDENCE_CARD_TAGS_LINE, EVIDENCE_CARD_TEMPLATE, EVIDENCE_CARDS_FOOTER,
    EVIDENCE_CARDS_HEADER, EXECUTIVE_SUMMARY_TEMPLATE, FINDINGS_TABLE_FOOTER,
    FINDINGS_TABLE_HEADER, FINDINGS_TABLE_ROW, HEADER_TEMPLATE, HTML_SKELETON,
    LIMITATIONS_ITEM, LIMITATIONS_NONE_TEMPLATE, LIMITATIONS_TEMPLATE_FOOTER,
    LIMITATIONS_TEMPLATE_HEADER, METHODOLOGY_PARAM_ROW, METHODOLOGY_REF_ROW,
    METHODOLOGY_TEMPLATE, REPORT_CSS, SECTION_ID_EVIDENCE_CARDS,
    SECTION_ID_EXECUTIVE_SUMMARY, SECTION_ID_FINDINGS_TABLE, SECTION_ID_HEADER,
    SECTION_ID_LIMITATIONS, SECTION_ID_METHODOLOGY, SECTION_ID_SIGN_OFF,
    SIGN_OFF_TEMPLATE)

# Empty findings table message when no variants are present.
_EMPTY_FINDINGS_MESSAGE = (
    "No clinically significant variants meeting triage criteria "
    "were identified within the requested panel target areas."
)


_TABLE_GRID = "Table Grid"


class ReportTemplateEngine:
    """Renders assembled report sections into HTML, PDF, or DOCX.

    HTML output is self-contained with all CSS inlined. PDF uses
    WeasyPrint for rendering. DOCX uses python-docx with proper
    Word styles.
    """

    def render_html(self, sections: ReportSections) -> str:
        """Produce self-contained HTML with inlined CSS.

        Parameters
        ----------
        sections : ReportSections
            All assembled report section data.

        Returns
        -------
        str
            Complete HTML document as a string.
        """
        body_parts: list[str] = []

        body_parts.append(self._render_header(sections))
        body_parts.append(self._render_disclaimer())
        body_parts.append(self._render_executive_summary(sections))
        body_parts.append(self._render_findings_table(sections))
        body_parts.append(self._render_evidence_cards(sections))
        body_parts.append(self._render_limitations(sections))
        body_parts.append(self._render_methodology(sections))
        body_parts.append(self._render_sign_off(sections))

        body = "\n".join(body_parts)

        # Escape user-provided strings to prevent HTML
        # injection (XSS). patient_id and panel_name come
        # from CLI input and must not be trusted.
        safe_patient_id = html.escape(sections.header.patient_id)

        html_content = HTML_SKELETON.format(
            patient_id=safe_patient_id,
            css=REPORT_CSS,
            body=body,
        )

        return html_content

    def render_pdf(self, sections: ReportSections, output_path: Path) -> Path:
        """Render HTML to PDF via WeasyPrint.

        The PDF preserves all text as selectable and searchable.

        Parameters
        ----------
        sections : ReportSections
            All assembled report section data.
        output_path : Path
            Destination file path for the PDF.

        Returns
        -------
        Path
            The output path where the PDF was written.

        Raises
        ------
        ImportError
            If WeasyPrint is not installed.
        """
        try:
            import weasyprint  # type: ignore[import-not-found]  # noqa: F401
        except ImportError as exc:
            raise ImportError(
                "PDF output requires the 'weasyprint' package. "
                "Install it with: pip install weasyprint"
            ) from exc

        html_content = self.render_html(sections)

        output_path.parent.mkdir(parents=True, exist_ok=True)

        html_doc = weasyprint.HTML(string=html_content)
        html_doc.write_pdf(str(output_path))

        return output_path

    def render_docx(self, sections: ReportSections, output_path: Path) -> Path:
        """Render to DOCX via python-docx.

        Uses proper Word styles (Heading 1, Heading 2, Normal,
        Table Grid) so users can apply institutional templates.

        Parameters
        ----------
        sections : ReportSections
            All assembled report section data.
        output_path : Path
            Destination file path for the DOCX.

        Returns
        -------
        Path
            The output path where the DOCX was written.

        Raises
        ------
        ImportError
            If python-docx is not installed.
        """
        try:
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E501
        except ImportError as exc:
            raise ImportError(
                "DOCX output requires the 'python-docx' package. "
                "Install it with: pip install python-docx"
            ) from exc

        doc = Document()

        self._docx_add_header(doc, sections.header)
        self._docx_add_executive_summary(doc, sections.executive_summary)
        self._docx_add_findings_table(doc, sections.findings_table, WD_TABLE_ALIGNMENT)
        self._docx_add_evidence_cards(doc, sections.evidence_cards)
        self._docx_add_limitations(doc, sections.limitations)
        self._docx_add_methodology(doc, sections.methodology)
        self._docx_add_sign_off(doc, sections.sign_off)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return output_path

    def _docx_add_header(
        self,
        doc: Any,
        header: HeaderData,
    ) -> None:
        """Add header section to DOCX document."""
        doc.add_heading("Clinical Variant Report", level=1)
        doc.add_paragraph(f"Patient ID: {header.patient_id}")
        doc.add_paragraph(f"Gene Panel: {header.panel_name}")
        doc.add_paragraph(f"Analysis Date: {header.analysis_date}")
        doc.add_paragraph(f"Pipeline Version: {header.pipeline_version}")

    def _docx_add_executive_summary(
        self,
        doc: Any,
        summary: ExecutiveSummaryData,
    ) -> None:
        """Add executive summary section to DOCX document."""
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(f"Total variants analyzed: {summary.total_variants_analyzed}")
        doc.add_paragraph(f"Variants passed filters: {summary.variants_passed_filters}")
        doc.add_paragraph(f"Pathogenic: {summary.pathogenic_count}")
        doc.add_paragraph(f"Likely Pathogenic: {summary.likely_pathogenic_count}")
        doc.add_paragraph(f"Variants of Uncertain Significance: {summary.vus_count}")

    def _docx_add_findings_table(
        self,
        doc: Any,
        findings_table: list[FindingsRow],
        wd_table_alignment: Any,
    ) -> None:
        """Add findings table section to DOCX document."""
        doc.add_heading("Findings Table", level=1)
        if not findings_table:
            doc.add_paragraph(_EMPTY_FINDINGS_MESSAGE)
            return

        table = doc.add_table(rows=1, cols=5)
        table.style = _TABLE_GRID
        table.alignment = wd_table_alignment.CENTER
        self._set_repeat_header_row(table)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Gene"
        hdr_cells[1].text = "Consequence"
        hdr_cells[2].text = "Classification"
        hdr_cells[3].text = "Composite Rank"
        hdr_cells[4].text = "Location"

        for row_data in findings_table:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data.gene_name or "Intergenic"
            row_cells[1].text = row_data.consequence
            row_cells[2].text = row_data.classification
            row_cells[3].text = (
                f"{row_data.composite_rank:.3f}"
                if row_data.composite_rank is not None
                else "N/A"
            )
            row_cells[4].text = f"{row_data.chromosome}:{row_data.position}"

    def _docx_add_evidence_cards(
        self,
        doc: Any,
        evidence_cards: list[EvidenceCardData],
    ) -> None:
        """Add evidence cards section to DOCX document."""
        doc.add_heading("Evidence Cards", level=1)
        for card in evidence_cards:
            gene_label = card.gene_name or "Intergenic"
            doc.add_heading(f"{gene_label}: {card.consequence}", level=2)

            if card.allele_frequency_formatted:
                doc.add_paragraph(
                    f"Allele Frequency: {card.allele_frequency_formatted}"
                )
            if card.predictor_scores_formatted:
                doc.add_paragraph(
                    "Predictor Scores: " + "; ".join(card.predictor_scores_formatted)
                )
            if card.clinvar_assertion:
                doc.add_paragraph(f"ClinVar: {card.clinvar_assertion}")
            if card.inheritance_pattern:
                doc.add_paragraph(f"Inheritance: {card.inheritance_pattern}")
            if card.evidence_tags_with_explanations:
                doc.add_paragraph(
                    "ACMG Criteria: " + "; ".join(card.evidence_tags_with_explanations)
                )

            doc.add_paragraph(card.narrative)

    def _docx_add_limitations(
        self,
        doc: Any,
        limitations: list[str],
    ) -> None:
        """Add limitations section to DOCX document."""
        doc.add_heading("Limitations", level=1)
        if limitations:
            for limitation in limitations:
                doc.add_paragraph(limitation, style="List Bullet")
        else:
            doc.add_paragraph(
                "No data source limitations were encountered during this analysis."
            )

    def _docx_add_methodology(
        self,
        doc: Any,
        methodology: MethodologyData,
    ) -> None:
        """Add methodology section to DOCX document."""
        doc.add_heading("Methodology", level=1)
        doc.add_paragraph(f"Pipeline Version: {methodology.pipeline_version}")
        doc.add_paragraph(f"Analysis Timestamp: {methodology.analysis_timestamp}")

        if methodology.reference_files:
            doc.add_heading("Reference Files", level=2)
            ref_table = doc.add_table(rows=1, cols=2)
            ref_table.style = _TABLE_GRID
            ref_hdr = ref_table.rows[0].cells
            ref_hdr[0].text = "File"
            ref_hdr[1].text = "SHA-256 Checksum"
            for path, checksum in methodology.reference_files.items():
                row_cells = ref_table.add_row().cells
                row_cells[0].text = path
                row_cells[1].text = checksum

        if methodology.classification_parameters:
            doc.add_heading("Classification Parameters", level=2)
            param_table = doc.add_table(rows=1, cols=2)
            param_table.style = _TABLE_GRID
            param_hdr = param_table.rows[0].cells
            param_hdr[0].text = "Parameter"
            param_hdr[1].text = "Value"
            for (
                param_name,
                param_value,
            ) in methodology.classification_parameters.items():
                row_cells = param_table.add_row().cells
                row_cells[0].text = param_name
                row_cells[1].text = param_value

    def _docx_add_sign_off(
        self,
        doc: Any,
        sign_off: SignOffData,
    ) -> None:
        """Add sign-off section to DOCX document."""
        doc.add_heading("Sign-off", level=1)
        doc.add_paragraph(f"Reviewer: {sign_off.reviewer_name_placeholder}")
        doc.add_paragraph(f"Date: {sign_off.review_date_placeholder}")
        doc.add_paragraph(
            f"Digital Signature: {sign_off.digital_signature_placeholder}"
        )

    def _render_header(self, sections: ReportSections) -> str:
        """Render the header section HTML."""
        return HEADER_TEMPLATE.format(
            section_id=SECTION_ID_HEADER,
            patient_id=html.escape(sections.header.patient_id),
            panel_name=html.escape(sections.header.panel_name),
            analysis_date=sections.header.analysis_date,
            pipeline_version=sections.header.pipeline_version,
        )

    def _render_disclaimer(self) -> str:
        """Render the computational-only disclaimer banner."""
        return (
            '<div class="disclaimer" style="'
            "background-color: #fff3cd; border: 1px solid #ffc107; "
            "border-radius: 4px; padding: 12px 16px; margin: 16px 0; "
            'font-size: 0.9em;">'
            "<strong>Computational Analysis Only</strong>: "
            "This report was generated by automated computational analysis "
            "using the ACMG/AMP 2015 variant classification framework "
            "(Richards et al., Genet Med 2015;17:405-424). "
            "All findings require review and confirmation by a qualified "
            "clinical geneticist before clinical action. "
            "This report does not constitute a medical diagnosis."
            "</div>"
        )

    def _render_executive_summary(self, sections: ReportSections) -> str:
        """Render the executive summary section HTML."""
        summary = sections.executive_summary
        return EXECUTIVE_SUMMARY_TEMPLATE.format(
            section_id=SECTION_ID_EXECUTIVE_SUMMARY,
            total_variants_analyzed=(summary.total_variants_analyzed),
            variants_passed_filters=(summary.variants_passed_filters),
            pathogenic_count=summary.pathogenic_count,
            likely_pathogenic_count=(summary.likely_pathogenic_count),
            vus_count=summary.vus_count,
        )

    def _render_findings_table(self, sections: ReportSections) -> str:
        """Render the findings table section HTML."""
        parts: list[str] = []
        parts.append(FINDINGS_TABLE_HEADER.format(section_id=SECTION_ID_FINDINGS_TABLE))

        if sections.findings_table:
            for row in sections.findings_table:
                rank_str = (
                    f"{row.composite_rank:.3f}"
                    if row.composite_rank is not None
                    else "N/A"
                )
                parts.append(
                    FINDINGS_TABLE_ROW.format(
                        gene_name=(row.gene_name or "Intergenic"),
                        consequence=row.consequence,
                        classification=row.classification,
                        composite_rank=rank_str,
                        chromosome=row.chromosome,
                        position=row.position,
                    )
                )
        else:
            parts.append(
                f'            <tr><td colspan="5">'
                f"{_EMPTY_FINDINGS_MESSAGE}</td></tr>\n"
            )

        parts.append(FINDINGS_TABLE_FOOTER)
        return "".join(parts)

    def _render_evidence_cards(self, sections: ReportSections) -> str:
        """Render the evidence cards section HTML."""
        parts: list[str] = []
        parts.append(EVIDENCE_CARDS_HEADER.format(section_id=SECTION_ID_EVIDENCE_CARDS))

        for card in sections.evidence_cards:
            details_parts: list[str] = []

            if card.allele_frequency_formatted:
                details_parts.append(
                    EVIDENCE_CARD_AF_LINE.format(
                        allele_frequency_formatted=(card.allele_frequency_formatted)
                    )
                )

            if card.predictor_scores_formatted:
                details_parts.append(
                    EVIDENCE_CARD_SCORES_LINE.format(
                        scores="; ".join(card.predictor_scores_formatted)
                    )
                )

            if card.clinvar_assertion:
                details_parts.append(
                    EVIDENCE_CARD_CLINVAR_LINE.format(
                        clinvar_assertion=card.clinvar_assertion
                    )
                )

            if card.inheritance_pattern:
                details_parts.append(
                    EVIDENCE_CARD_INHERITANCE_LINE.format(
                        inheritance_pattern=(card.inheritance_pattern)
                    )
                )

            if card.evidence_tags_with_explanations:
                details_parts.append(
                    EVIDENCE_CARD_TAGS_LINE.format(
                        tags="; ".join(card.evidence_tags_with_explanations)
                    )
                )

            gene_label = card.gene_name or "Intergenic"
            parts.append(
                EVIDENCE_CARD_TEMPLATE.format(
                    gene_name=gene_label,
                    consequence=card.consequence,
                    details="".join(details_parts),
                    narrative=card.narrative,
                )
            )

        parts.append(EVIDENCE_CARDS_FOOTER)
        return "".join(parts)

    def _render_limitations(self, sections: ReportSections) -> str:
        """Render the limitations section HTML."""
        if not sections.limitations:
            return LIMITATIONS_NONE_TEMPLATE.format(section_id=SECTION_ID_LIMITATIONS)

        parts: list[str] = []
        parts.append(
            LIMITATIONS_TEMPLATE_HEADER.format(section_id=SECTION_ID_LIMITATIONS)
        )
        for limitation in sections.limitations:
            parts.append(LIMITATIONS_ITEM.format(limitation=limitation))
        parts.append(LIMITATIONS_TEMPLATE_FOOTER)
        return "".join(parts)

    def _render_methodology(self, sections: ReportSections) -> str:
        """Render the methodology section HTML."""
        methodology = sections.methodology

        ref_rows = ""
        for path, checksum in methodology.reference_files.items():
            ref_rows += METHODOLOGY_REF_ROW.format(path=path, checksum=checksum)

        param_rows = ""
        for param_name, param_value in methodology.classification_parameters.items():
            param_rows += METHODOLOGY_PARAM_ROW.format(
                param_name=param_name,
                param_value=param_value,
            )

        return METHODOLOGY_TEMPLATE.format(
            section_id=SECTION_ID_METHODOLOGY,
            pipeline_version=methodology.pipeline_version,
            analysis_timestamp=methodology.analysis_timestamp,
            reference_rows=ref_rows,
            parameter_rows=param_rows,
        )

    def _render_sign_off(self, sections: ReportSections) -> str:
        """Render the sign-off section HTML."""
        sign_off = sections.sign_off
        return SIGN_OFF_TEMPLATE.format(
            section_id=SECTION_ID_SIGN_OFF,
            reviewer_name=(sign_off.reviewer_name_placeholder),
            review_date=sign_off.review_date_placeholder,
            digital_signature=(sign_off.digital_signature_placeholder),
        )

    @staticmethod
    def _set_repeat_header_row(table: Any) -> None:
        """Enable 'repeat header row' on a DOCX table.

        This makes column headers appear on every page when the
        table spans multiple pages. Uses python-docx private API;
        fails silently if the internal structure changes.
        """
        try:
            from docx.oxml.ns import qn

            tbl_pr = table.rows[0]._tr
            tr_pr = tbl_pr.get_or_add_trPr()
            repeat_header = tr_pr.makeelement(qn("w:tblHeader"), {})
            tr_pr.append(repeat_header)
        except (AttributeError, ImportError):
            pass
